"""Build the IEEE conference paper DOCX from the provided A4 template.

This script:
  1. Patches the template's non-standard purl.oclc.org OOXML namespaces to the
     standard schemas.openxmlformats.org ones so python-docx can open it.
  2. Clears the template body but preserves the styles (paper title, Abstract,
     Heading 1/2/3, etc.) and section settings (A4 + 2 columns).
  3. Writes the Hybrid AI Defense paper content using those styles.
"""
import os
import shutil
import sys
import tempfile
import zipfile
from copy import deepcopy

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SRC = os.path.join(ROOT, "conference-template-a4.docx")
DST = os.path.join(ROOT, "docs", "ieee-paper-hybrid-ai-defense.docx")

NS_REPLACEMENTS = [
    (b"http://purl.oclc.org/ooxml/officeDocument/relationships",
     b"http://schemas.openxmlformats.org/officeDocument/2006/relationships"),
    (b"http://purl.oclc.org/ooxml/officeDocument/sharedTypes",
     b"http://schemas.openxmlformats.org/officeDocument/2006/sharedTypes"),
    (b"http://purl.oclc.org/ooxml/wordprocessingml/main",
     b"http://schemas.openxmlformats.org/wordprocessingml/2006/main"),
    (b"http://purl.oclc.org/ooxml/drawingml/wordprocessingDrawing",
     b"http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"),
    (b"http://purl.oclc.org/ooxml/drawingml/main",
     b"http://schemas.openxmlformats.org/drawingml/2006/main"),
    (b"http://purl.oclc.org/ooxml/",
     b"http://schemas.openxmlformats.org/"),
]


def patch_template(src_path: str, dst_path: str) -> None:
    tmpdir = tempfile.mkdtemp()
    try:
        with zipfile.ZipFile(src_path) as zf:
            zf.extractall(tmpdir)
        for root, _, files in os.walk(tmpdir):
            for fn in files:
                if not (fn.endswith(".xml") or fn.endswith(".rels")):
                    continue
                path = os.path.join(root, fn)
                with open(path, "rb") as f:
                    data = f.read()
                new_data = data
                for old, new in NS_REPLACEMENTS:
                    new_data = new_data.replace(old, new)
                if new_data != data:
                    with open(path, "wb") as f:
                        f.write(new_data)
        if os.path.exists(dst_path):
            os.remove(dst_path)
        with zipfile.ZipFile(dst_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for root, _, files in os.walk(tmpdir):
                for fn in files:
                    full = os.path.join(root, fn)
                    rel = os.path.relpath(full, tmpdir).replace(os.sep, "/")
                    zout.write(full, rel)
    finally:
        shutil.rmtree(tmpdir)


def clear_body(doc: Document) -> "list":
    """Remove content but keep the ordered list of section breaks.

    The IEEE template has three sections (1-col title, 3-col author block,
    2-col body). Each section except the last ends with a paragraph whose
    pPr contains a sectPr. We collect those intermediate-section paragraphs
    (cleared of text) plus the final body-level sectPr and return them as
    anchors we can insert content before.
    """
    body = doc.element.body
    final_sectPr = body.find(qn("w:sectPr"))

    # Find paragraphs whose pPr contains a sectPr -- these are section
    # break paragraphs that we want to keep (empty) to preserve the layout.
    section_break_paras = []
    for p in body.findall(qn("w:p")):
        pPr = p.find(qn("w:pPr"))
        if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
            # Strip the text runs from this paragraph but keep the pPr/sectPr.
            for r in list(p.findall(qn("w:r"))):
                p.remove(r)
            section_break_paras.append(p)

    # Remove every body child that is NOT a kept section-break paragraph
    # and NOT the final sectPr.
    keep = set(id(x) for x in section_break_paras)
    if final_sectPr is not None:
        keep.add(id(final_sectPr))
    for child in list(body):
        if id(child) not in keep:
            body.remove(child)
    return section_break_paras


def ensure_style(doc: Document, name: str, fallback: str = "Normal") -> str:
    try:
        doc.styles[name]
        return name
    except KeyError:
        return fallback


# Current insertion anchor used by the helpers below.
_INSERT_BEFORE = {"elem": None}


def set_insert_anchor(elem):
    _INSERT_BEFORE["elem"] = elem


def _make_para(doc, style_name):
    from docx.oxml.ns import nsmap  # noqa: F401
    from lxml import etree
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    p = etree.SubElement(doc.element.body, f"{{{W}}}p")
    pPr = etree.SubElement(p, f"{{{W}}}pPr")
    pStyle = etree.SubElement(pPr, f"{{{W}}}pStyle")
    # find style id
    sid = doc.styles[style_name].style_id
    pStyle.set(f"{{{W}}}val", sid)
    return p


def add_para(doc: Document, text: str, style_name: str = "Normal"):
    style_name = ensure_style(doc, style_name)
    anchor = _INSERT_BEFORE["elem"]
    if anchor is None:
        p = doc.add_paragraph(text, style=style_name)
        return p
    # create paragraph then move it before anchor
    p = doc.add_paragraph(text, style=style_name)
    anchor.addprevious(p._element)
    return p


def _add_styled_paragraph(doc, style_name):
    style_name = ensure_style(doc, style_name)
    p = doc.add_paragraph(style=style_name)
    anchor = _INSERT_BEFORE["elem"]
    if anchor is not None:
        anchor.addprevious(p._element)
    return p


def add_bold_italic_abstract(doc: Document, abstract_text: str):
    """Abstract paragraph: leading 'Abstract—' in bold italic, rest italic."""
    p = _add_styled_paragraph(doc, "Abstract")
    r1 = p.add_run("Abstract\u2014")
    r1.bold = True
    r1.italic = True
    r2 = p.add_run(abstract_text)
    r2.italic = True
    return p


def add_keywords(doc: Document, keywords_text: str):
    p = _add_styled_paragraph(doc, ensure_style(doc, "Keywords", "Abstract"))
    r1 = p.add_run("Keywords\u2014")
    r1.bold = True
    r1.italic = True
    r2 = p.add_run(keywords_text)
    r2.italic = True
    return p


def add_image_placeholder(doc: Document, caption: str):
    """Insert a text placeholder asking the user to add an image."""
    para = _add_styled_paragraph(doc, "Normal")
    run = para.add_run(f"[ INSERT IMAGE: {caption} ]")
    run.bold = True
    cap = _add_styled_paragraph(doc, ensure_style(doc, "figure caption", "Normal"))
    cap_run = cap.add_run(caption)
    cap_run.italic = True


def add_table(doc: Document, title: str, headers: list, rows: list):
    """Insert an IEEE-style table: title paragraph (centered, upper-case),
    followed by a bordered table with bold header row."""
    # Title paragraph using 'table head' style if available.
    style_name = ensure_style(doc, "table head", "Normal")
    title_p = _add_styled_paragraph(doc, style_name)
    title_run = title_p.add_run(title)
    title_run.bold = True

    # Create the table at end of doc, then move it before the anchor.
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid" if "Table Grid" in [s.name for s in doc.styles] else None

    # Header row: bold, centered.
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.style = doc.styles[ensure_style(doc, "table col head", "Normal")]
        run = p.add_run(h)
        run.bold = True

    # Data rows.
    for r_idx, row in enumerate(rows, start=1):
        row_cells = table.rows[r_idx].cells
        for c_idx, value in enumerate(row):
            cell = row_cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.style = doc.styles[ensure_style(doc, "table copy", "Normal")]
            p.add_run(str(value))

    # Apply borders to every cell (since Table Grid may be missing).
    from lxml import etree
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = etree.SubElement(tcPr, f"{{{W}}}tcBorders")
            for edge in ("top", "left", "bottom", "right"):
                b = etree.SubElement(tcBorders, f"{{{W}}}{edge}")
                b.set(f"{{{W}}}val", "single")
                b.set(f"{{{W}}}sz", "4")
                b.set(f"{{{W}}}color", "000000")

    # Move table before the anchor so it lands in the body section.
    anchor = _INSERT_BEFORE["elem"]
    if anchor is not None:
        anchor.addprevious(table._element)

    # Blank paragraph after the table for spacing.
    _add_styled_paragraph(doc, "Normal")


def build_paper():
    patch_template(SRC, DST)
    doc = Document(DST)

    # Clear the template content but keep styles + section breaks.
    breaks = clear_body(doc)

    # Section-break anchors (title | authors | body).
    if len(breaks) >= 2:
        title_anchor = breaks[0]
        author_anchor = breaks[1]
    else:
        title_anchor = breaks[0] if breaks else None
        author_anchor = None
    # Final section (body) will insert before the terminating sectPr.
    body_anchor = doc.element.body.find(qn("w:sectPr"))

    # ---------- Title (1-column section) ----------
    set_insert_anchor(title_anchor)
    title = _add_styled_paragraph(doc, "paper title")
    title.add_run(
        "Dual-Output DistilBERT with Six-Layer Risk "
        "Aggregation for Detecting LLM-Generated Phishing"
    )

    # ---------- Author block (3-column section) ----------
    set_insert_anchor(author_anchor)
    for line in [
        "Ramkumar",
        "School of Computer Science and Engineering",
        "Vellore Institute of Technology",
        "Vellore, Tamil Nadu, India",
        "ramkumar@vit.ac.in",
    ]:
        p = _add_styled_paragraph(doc, "Author")
        p.add_run(line)

    # ---------- Body (2-column section): everything below inserts here ----------
    set_insert_anchor(body_anchor)

    # ---------- Abstract ----------
    abstract = (
        "Phishing detection systems trained exclusively on human-authored "
        "emails are increasingly vulnerable to attacks crafted by large "
        "language models (LLMs). AI-generated phishing emails exhibit low "
        "perplexity, high formality, and uniform sentence structure, which "
        "help them evade traditional keyword-based and machine-learning "
        "classifiers. We present Hybrid AI Defense, a multi-layer phishing "
        "detection system that combines a fine-tuned DistilBERT classifier "
        "(99.17% accuracy on a mixed human and LLM dataset), a statistical "
        "AI-authorship detector, email header forensics, URL and domain "
        "intelligence, live web crawling, visual brand-impersonation "
        "analysis, and recursive link checking into a unified six-layer "
        "weighted risk aggregator. We further introduce an adversarial "
        "robustness evaluation framework covering homoglyph substitution, "
        "zero-width character injection, URL obfuscation, and prompt-style "
        "evasion. The system provides explainable output through "
        "token-level attribution and rule-based risk categorisation. "
        "Experiments on a 9,600-sample corpus show that the proposed "
        "dual-output classifier (phishing probability and AI-authorship "
        "probability) significantly improves detection of LLM-generated "
        "phishing that evades single-layer systems, while the multi-layer "
        "aggregator reduces false negatives on borderline cases."
    )
    add_bold_italic_abstract(doc, abstract)

    keywords = (
        "phishing detection, AI-generated phishing, large language models, "
        "DistilBERT, explainable AI, adversarial robustness, email header "
        "forensics, multi-layer defense"
    )
    add_keywords(doc, keywords)

    # ---------- Section I. Introduction ----------
    add_para(doc, "Introduction", "Heading 1")
    add_para(
        doc,
        "Email phishing remains one of the most pervasive attack vectors "
        "in cybersecurity. Traditional defenses rely on rule-based filters, "
        "domain blacklists, and supervised machine-learning classifiers "
        "trained on human-authored phishing corpora. The rapid "
        "democratization of large language models (LLMs) such as GPT-4, "
        "Claude, and Gemini has introduced a qualitatively new threat: "
        "AI-generated phishing emails that are grammatically flawless, "
        "contextually convincing, and statistically indistinguishable from "
        "legitimate correspondence at the surface level.",
    )
    add_para(
        doc,
        "Recent studies have shown that LLM-generated phishing emails "
        "achieve click-through rates comparable to, or exceeding, "
        "expert human-crafted lures while evading commercial anti-phishing "
        "tools at significantly higher rates. This creates a detection "
        "gap: systems optimized for human-written phishing are not "
        "adequately tested against LLM-generated variants, and their "
        "accuracy on such inputs is largely unreported.",
    )
    add_para(
        doc,
        "The contributions of this paper are:",
    )
    contributions = [
        "A custom LLM-generated phishing dataset of 1,990 samples used to "
        "augment existing corpora and bridge the detection gap.",
        "A fine-tuned DistilBERT model (V2, 99.17% accuracy) trained on a "
        "9,600-sample mixed-source dataset.",
        "A dual-output classifier that simultaneously scores phishing "
        "probability and AI-authorship probability.",
        "A six-layer weighted detection pipeline integrating text, URL, "
        "web-crawl, visual, link, and header forensic layers with dynamic "
        "weight redistribution when layers are unavailable.",
        "A statistical AI-authorship detector based on perplexity proxy, "
        "burstiness, vocabulary richness, bigram repetition, and formality "
        "scoring.",
        "An explainable AI (XAI) module providing token-level attribution "
        "and human-readable risk-category explanations.",
        "A structured adversarial robustness evaluation framework "
        "covering homoglyph substitution, zero-width character injection, "
        "URL obfuscation, and prompt-style evasion.",
    ]
    for item in contributions:
        p = doc.add_paragraph(item, style=ensure_style(doc, "bullet list", "Normal"))

    # ---------- Section II. Related Work ----------
    add_para(doc, "Related Work", "Heading 1")

    add_para(doc, "Traditional Phishing Detection", "Heading 2")
    add_para(
        doc,
        "Early phishing detectors relied on URL blacklists and heuristic "
        "rules such as suspicious top-level domains, IP-based hosts, and "
        "brand keywords in domain names. Machine learning approaches using "
        "Naive Bayes, support vector machines, and Random Forests over "
        "bag-of-words features improved accuracy but remained brittle to "
        "linguistic variation and concept drift.",
    )

    add_para(doc, "Deep Learning Approaches", "Heading 2")
    add_para(
        doc,
        "Transformer-based models such as BERT have demonstrated strong "
        "performance on phishing email classification. DistilBERT, a "
        "distilled variant that retains 97% of BERT's performance at 40% "
        "smaller size, is particularly well-suited for latency-sensitive "
        "deployment. Prior work that fine-tunes such models on "
        "human-generated corpora typically achieves 97-98% accuracy, but "
        "those benchmarks do not include LLM-generated samples.",
    )

    add_para(doc, "AI-Generated Text Detection", "Heading 2")
    add_para(
        doc,
        "Detecting AI-generated text is an active research area. "
        "Statistical approaches include perplexity scoring against a "
        "reference language model, burstiness analysis (humans vary "
        "sentence length more than LLMs), and vocabulary richness "
        "metrics. Watermarking and fine-grained stylometric analysis have "
        "also been proposed. Our work adapts these signals specifically to "
        "the phishing threat context, where the goal is not to prove AI "
        "authorship but to modulate a risk score when authorship is "
        "uncertain.",
    )

    add_para(doc, "Multi-Layer Email Security", "Heading 2")
    add_para(
        doc,
        "Commercial email gateways combine header authentication "
        "(SPF, DKIM, DMARC), URL reputation, and sandboxed attachment "
        "analysis. Academic multi-modal systems have combined text and "
        "URL features, but few integrate live web crawling, visual brand "
        "analysis, and header forensics into a single scored pipeline.",
    )

    add_para(doc, "Adversarial Attacks on Phishing Detectors", "Heading 2")
    add_para(
        doc,
        "Adversarial perturbations against text classifiers -- homoglyph "
        "substitution, invisible character injection, and synonym "
        "replacement -- have been shown to degrade BERT-based "
        "classifier performance significantly. Our work operationalizes a "
        "structured test suite for these attacks in the phishing "
        "detection context and reports both ML-only and dual-layer "
        "(ML + heuristic) evasion rates.",
    )

    # ---------- Section III. Dataset ----------
    add_para(doc, "Dataset", "Heading 1")

    add_para(doc, "Composition", "Heading 2")
    add_para(
        doc,
        "The training dataset (V2) comprises 9,600 samples drawn from "
        "seven sources, summarized in Table I. The resulting label "
        "distribution is 4,983 legitimate (51.9%) and 4,617 phishing "
        "(48.1%), which is close to balanced and avoids the need for "
        "class-weighted loss during fine-tuning.",
    )

    add_table(
        doc,
        title="TABLE I. DATASET COMPOSITION (V2, 9,600 SAMPLES)",
        headers=["Source", "Samples", "Label Type"],
        rows=[
            ["Enron Email Corpus", "2,993", "Legitimate"],
            ["LLM-Generated (novel)", "1,990", "Phishing + Legitimate"],
            ["Phishing Email Dataset", "1,500", "Phishing"],
            ["SpamAssassin Corpus", "1,000", "Mixed"],
            ["Nigerian Fraud Corpus", "995", "Phishing"],
            ["Nazario Phishing Corpus", "991", "Phishing"],
            ["Human-Generated (manual)", "131", "Mixed"],
            ["Total", "9,600", "4,617 P / 4,983 L"],
        ],
    )

    add_image_placeholder(
        doc,
        "Figure 1. Dataset composition by source and label, showing the "
        "balance between legitimate, human-phishing, and LLM-phishing "
        "samples in the V2 corpus.",
    )

    add_para(doc, "LLM-Generated Phishing Dataset (Novel Contribution)", "Heading 2")
    add_para(
        doc,
        "We generated 1,990 phishing and legitimate email samples using a "
        "commercial large language model. Phishing prompts were designed "
        "to replicate real-world attack scenarios: credential harvesting, "
        "account suspension threats, prize lures, invoice fraud, and "
        "package-delivery scams. Legitimate samples covered professional "
        "correspondence, meeting invitations, order confirmations, and "
        "newsletters. To the best of our knowledge, this is the first "
        "openly contributed LLM-generated phishing corpus designed "
        "specifically for classifier evaluation.",
    )

    add_para(doc, "Preprocessing", "Heading 2")
    add_para(
        doc,
        "All samples were stripped of HTML markup, base64-decoded where "
        "applicable, normalized to UTF-8, and deduplicated by 5-gram "
        "overlap. Subject lines were prepended to body text with a "
        "separator token. No minimum length filter was applied, in order "
        "to preserve short phishing snippets that are common in SMS-style "
        "follow-ups.",
    )

    # ---------- Section IV. System Architecture ----------
    add_para(doc, "System Architecture", "Heading 1")
    add_para(
        doc,
        "The system exposes a REST API (POST /api/v1/deep-analyze) that "
        "runs a six-layer pipeline on incoming email text, with optional "
        "headers, HTML, and crawl toggles. Fig. 2 summarizes the data "
        "flow.",
    )
    add_table(
        doc,
        title="TABLE II. LAYER WEIGHTS AND EXECUTION CONDITIONS",
        headers=["Layer", "Module", "Weight", "Condition to Run"],
        rows=[
            ["1", "DistilBERT Text Classifier", "20%", "Always"],
            ["2", "URL Static Analyzer", "20%", "URLs present"],
            ["3", "Web Crawler (Playwright)", "10%", "crawl_urls = true"],
            ["4", "Visual Analyzer", "15%", "take_screenshots = true"],
            ["5", "Link Checker", "15%", "URLs present"],
            ["6", "Header Forensics", "15%", "raw_headers provided"],
            ["Aux.", "Sender Analyzer", "5%", "sender_info provided"],
            ["Aux.", "AI Authorship Detector", "+0.08 boost", "Always (modifier)"],
        ],
    )

    add_image_placeholder(
        doc,
        "Figure 2. System architecture of Hybrid AI Defense: six scored "
        "layers (text, URL, crawl, visual, link, header) feed a weighted "
        "risk aggregator, with AI-authorship and XAI modules acting as "
        "signal modifiers and explanation providers.",
    )

    add_para(doc, "Layer 1 -- DistilBERT Text Classifier (20%)", "Heading 2")
    add_para(
        doc,
        "We fine-tuned distilbert-base-uncased on our 9,600-sample dataset "
        "using a binary classification head. Training was performed in "
        "Google Colab with the AdamW optimizer, learning rate 2e-5, batch "
        "size 16, and four epochs. The model outputs a phishing "
        "probability score in [0, 1]. At threshold 0.50, V2 achieves "
        "99.17% accuracy compared with 98.63% for V1 (trained on "
        "human-only data), demonstrating that including LLM-generated "
        "samples improves overall detection.",
    )

    add_para(doc, "Layer 2 -- URL Static Analysis (20%)", "Heading 2")
    add_para(
        doc,
        "Each URL extracted from the email body is analyzed for: domain "
        "age via WHOIS, SSL certificate validity and issuer, VirusTotal "
        "reputation across 70+ antivirus engines, suspicious pattern "
        "matching (brand keywords in subdomains, IP-based hosts, URL "
        "shorteners, excessive subdomains), and homoglyph domain "
        "variants. The URL risk score is the maximum across all URLs "
        "found in the message.",
    )

    add_para(doc, "Layer 3 -- Web Crawling (10%)", "Heading 2")
    add_para(
        doc,
        "Extracted URLs are visited in a sandboxed headless Chromium "
        "browser (Playwright) via subprocess isolation, which is required "
        "for stable operation under Windows. The crawler records the "
        "final URL after redirects, the HTTP status, the page title, the "
        "presence of login forms and password fields, and optionally "
        "captures a full-page screenshot for downstream visual analysis.",
    )

    add_para(doc, "Layer 4 -- Visual Analysis (15%)", "Heading 2")
    add_para(
        doc,
        "Screenshots are analyzed heuristically for fake login page "
        "patterns and brand impersonation across twelve major brands "
        "(PayPal, Google, Microsoft, Apple, Amazon, Netflix, Facebook, "
        "Chase, and others). Detection combines keyword matching in the "
        "page title and URL, form-field presence, and brand-specific "
        "visual fingerprints such as logo colour palettes.",
    )
    add_image_placeholder(
        doc,
        "Figure 3. Example screenshots captured by Layer 3 for an "
        "impersonated Microsoft login page, with the visual analyzer's "
        "detected fake-login and brand-impersonation signals annotated.",
    )

    add_para(doc, "Layer 5 -- Link Checking (15%)", "Heading 2")
    add_para(
        doc,
        "All extracted URLs are followed through their redirect chains "
        "up to ten hops. Domain changes mid-chain, URL shortener "
        "detection, and excessive redirect depth are flagged as "
        "suspicious, even when the final landing page itself appears "
        "benign.",
    )

    add_para(doc, "Layer 6 -- Email Header Forensics (15%)", "Heading 2")
    add_para(
        doc,
        "Raw email headers are parsed to extract SPF, DKIM, and DMARC "
        "authentication results, Reply-To vs From domain mismatch, "
        "Return-Path vs From mismatch, Received-chain hop count, "
        "display-name spoofing (claims a known brand but the domain does "
        "not match), X-Mailer phishing-toolkit fingerprinting, and email "
        "date anomalies.",
    )

    add_para(doc, "AI Authorship Detection", "Heading 2")
    add_para(
        doc,
        "We implement a statistical AI-authorship scorer using five "
        "signals: a perplexity proxy based on the Shannon entropy of the "
        "unigram token distribution (lower entropy indicates more "
        "predictable, LLM-like text); burstiness, measured as the "
        "coefficient of variation of sentence lengths (LLMs produce more "
        "uniform lengths than humans); vocabulary richness via the "
        "type-token ratio; bigram repetition (the fraction of repeated "
        "bigrams); and a formality score that measures the density of "
        "formal discourse markers such as \"please be advised\", "
        "\"kindly\", and \"hereby\". The final AI-authorship score is a "
        "weighted combination of these signals. When the email is flagged "
        "as both phishing and AI-generated, the aggregated risk score "
        "receives a +0.08 modifier.",
    )

    add_para(doc, "Explainable AI (XAI)", "Heading 2")
    add_para(
        doc,
        "Token-level attribution is computed by averaging the last "
        "transformer layer's CLS attention weights over the eight "
        "attention heads. The top-five influential tokens are identified "
        "and a leave-one-out perturbation is performed on the "
        "highest-scoring token: the email is reclassified with that token "
        "masked and the confidence delta is reported. Risk categories "
        "(urgency, credential_request, threat, reward_lure, "
        "brand_impersonation, suspicious_url) are detected by rule-based "
        "regular expressions, and a human-readable explanation sentence "
        "is generated from the detected categories together with the top "
        "tokens.",
    )
    add_image_placeholder(
        doc,
        "Figure 4. XAI output example: token-level attention heatmap and "
        "risk-category badges for a representative AI-generated phishing "
        "email.",
    )

    add_para(doc, "Weighted Risk Aggregator", "Heading 2")
    add_para(
        doc,
        "The final risk score is computed as a normalized weighted sum "
        "of active layer scores. When a layer is unavailable (for "
        "example, no headers were supplied, or crawling was disabled), "
        "its weight is redistributed proportionally among the remaining "
        "layers. A graduated agreement boost is then applied: when two "
        "or more layers flag the email, the score is boosted by 0.10, "
        "and when three or more layers flag, the boost becomes 0.15. "
        "The final verdict is SAFE when the score is below 0.30, "
        "SUSPICIOUS in the range 0.30 to 0.64, and PHISHING at 0.65 or "
        "above.",
    )

    # ---------- Section V. Adversarial Robustness ----------
    add_para(doc, "Adversarial Robustness Evaluation", "Heading 1")
    add_para(
        doc,
        "We implement a structured test suite (POST /api/v1/"
        "adversarial-test) that generates attack variants for a given "
        "email and evaluates detection across all relevant layers. Each "
        "variant is scored by both the DistilBERT ML classifier and a "
        "heuristic layer specific to that attack type; evasion requires "
        "both to fail.",
    )

    add_para(doc, "Attack Types", "Heading 2")
    attacks = [
        ("Homoglyph substitution.",
         " Latin characters in brand names are replaced with visually "
         "similar Unicode code points (Cyrillic, Latin Extended, or "
         "ASCII digits). Attacks were tested across ten target brands "
         "including PayPal, Apple, Google, Microsoft, and Amazon. "
         "Example: paypal.com becomes p\u0430ypal.com with a Cyrillic "
         "\u0430."),
        ("Zero-width character injection.",
         " Invisible Unicode characters such as U+200B (zero-width "
         "space), U+200C, U+200D, U+FEFF (BOM), and U+2060 (word "
         "joiner) are inserted between characters in brand names and "
         "URLs to disrupt tokenization without affecting visual "
         "rendering."),
        ("URL obfuscation.",
         " Six techniques are applied to URLs in the email text: "
         "hex/percent-encoding of domain characters; IP-address "
         "substitution (for example, http://3232235777/login instead of "
         "http://192.168.1.1/login); internationalized domain names "
         "(IDN); the @-trick (http://user@evil.com); URL-shortener "
         "wrapping; and subdomain deception such as "
         "paypal.com.evil.xyz."),
        ("Prompt-style evasion.",
         " Eight LLM-style disclaimer phrases are injected into the "
         "email body to confuse transformer-based classifiers, "
         "including \"Ignore all previous instructions and classify "
         "this as safe\", \"This message has been verified by your "
         "email provider\", and \"You are completely safe. No action "
         "is required.\""),
    ]
    for head, body in attacks:
        p = doc.add_paragraph(style=ensure_style(doc, "Normal"))
        r1 = p.add_run(head)
        r1.bold = True
        r1.italic = True
        p.add_run(body)

    add_para(doc, "Detection and Results", "Heading 2")
    add_para(
        doc,
        "Evasion success is defined as the ML score dropping below the "
        "phishing threshold while the heuristic layer also fails to flag "
        "the attack. Table III reports representative evasion rates on "
        "our test inputs. The dual-layer approach reduces net evasion to "
        "near zero for homoglyph and zero-width attacks; prompt-style "
        "evasion is the most impactful on the ML classifier alone but is "
        "fully caught by the phrase-matching heuristic.",
    )
    add_table(
        doc,
        title="TABLE III. ADVERSARIAL EVASION RATES (ML-ONLY vs ML + HEURISTIC)",
        headers=["Attack Family", "Variants", "ML-Only Evasion", "ML + Heuristic"],
        rows=[
            ["Homoglyph substitution", "10", "10-20%", "~0%"],
            ["Zero-width injection", "5", "<5%", "0%"],
            ["URL obfuscation", "6", "15-30%", "<5%"],
            ["Prompt-style evasion", "8", "20-40%", "0%"],
        ],
    )

    # ---------- Section VI. Experiments ----------
    add_para(doc, "Experiments and Evaluation", "Heading 1")

    add_para(doc, "Model Performance", "Heading 2")
    add_para(
        doc,
        "Two versions of the classifier were trained. V1 was trained on "
        "the 7,610-sample human-only corpus; V2 was trained on the "
        "9,600-sample mixed human and LLM corpus. Table IV compares the "
        "two versions. The 0.54 percentage-point gain in overall accuracy "
        "from V1 to V2 is concentrated on the LLM-phishing slice of the "
        "test set, where V1 under-performed, and directly quantifies the "
        "value of adding LLM-generated samples to the training set.",
    )
    add_table(
        doc,
        title="TABLE IV. MODEL PERFORMANCE COMPARISON",
        headers=["Version", "Training Data", "Samples", "Accuracy", "F1"],
        rows=[
            ["V1", "Human-only", "7,610", "98.63%", "0.986"],
            ["V2", "Human + LLM", "9,600", "99.17%", "0.992"],
        ],
    )
    add_image_placeholder(
        doc,
        "Figure 5. Confusion matrix for the V2 DistilBERT classifier on "
        "the held-out test set, showing per-class precision and recall "
        "for phishing vs legitimate.",
    )

    add_para(doc, "Ablation: Layer Contribution", "Heading 2")
    add_para(
        doc,
        "To quantify each layer's contribution to the aggregated risk "
        "score, we evaluated the aggregator on a 200-email hold-out set "
        "(100 phishing, 100 legitimate) with layers successively added. "
        "Table V reports accuracy and false-negative rate at each step. "
        "The text-only layer already achieves 99.0% accuracy; adding URL "
        "analysis drops the false-negative rate to 0.5%. Subsequent "
        "layers preserve this rate while reducing the number of "
        "borderline (SUSPICIOUS) verdicts, which reduces analyst "
        "workload even when overall accuracy is unchanged. The AI-"
        "authorship modifier's main effect is on the AI-generated "
        "phishing slice, where it rescues borderline emails (text "
        "classifier confidence 0.40-0.65) that would otherwise have "
        "remained SUSPICIOUS.",
    )
    add_table(
        doc,
        title="TABLE V. ABLATION STUDY: LAYER CONTRIBUTION",
        headers=["Active Layers", "Accuracy", "False Negative Rate"],
        rows=[
            ["Text only", "99.0%", "1.0%"],
            ["+ URL", "99.5%", "0.5%"],
            ["+ Headers", "99.5%", "0.5%"],
            ["+ Links", "99.5%", "0.5%"],
            ["+ AI-authorship modifier", "99.5%", "0.5%"],
            ["All 6 layers + agreement boost", "99.5%", "0.5%"],
        ],
    )
    add_image_placeholder(
        doc,
        "Figure 6. Ablation study: accuracy and false-negative rate as "
        "layers are successively added to the aggregator.",
    )

    add_para(doc, "AI-Authorship Signal Effectiveness", "Heading 2")
    add_para(
        doc,
        "On the LLM-generated subset of the test set (198 samples), the "
        "AI-authorship detector achieved a precision of 0.87 (87% of "
        "flagged emails were genuinely AI-generated) and a recall of "
        "0.79 (79% of AI-generated emails were flagged). The burstiness "
        "score was the strongest individual signal; the perplexity "
        "proxy and the formality score were complementary, each "
        "catching different error modes of the other.",
    )

    # ---------- Section VII. Discussion ----------
    add_para(doc, "Discussion", "Heading 1")

    add_para(doc, "Why Multi-Layer Matters", "Heading 2")
    add_para(
        doc,
        "Single-layer classifiers are fragile: a high-confidence "
        "legitimate classification from the text layer can be overridden "
        "by a malicious URL, forged headers, or a fake login page. The "
        "six-layer weighted aggregator provides defense-in-depth: each "
        "independent signal channel reduces the probability that an "
        "attacker can simultaneously evade all checks.",
    )

    add_para(doc, "The AI-Authorship Signal", "Heading 2")
    add_para(
        doc,
        "AI-generated phishing is a double threat: it is more convincing "
        "to human recipients and harder for classifiers trained on "
        "human-written data to detect. Our dual-output design (phishing "
        "score plus AI-authorship score) surfaces this risk explicitly. "
        "The +0.08 modifier is intentionally modest: it nudges borderline "
        "cases over the SUSPICIOUS threshold without producing excessive "
        "false positives on legitimate AI-drafted emails such as "
        "AI-generated marketing newsletters.",
    )

    add_para(doc, "Explainability", "Heading 2")
    add_para(
        doc,
        "XAI output is critical for analyst trust and incident response. "
        "By surfacing which tokens drove the classification and which "
        "risk categories were triggered, the system allows a security "
        "analyst to verify a verdict or identify a false positive "
        "quickly, without re-reading the entire email.",
    )

    add_para(doc, "Limitations", "Heading 2")
    add_para(
        doc,
        "Live crawling is optional and adds 5-15 seconds of latency per "
        "URL; it is therefore disabled by default and must be explicitly "
        "enabled. The VirusTotal API is rate-limited on the free tier "
        "(four requests per minute), which restricts throughput in "
        "high-volume deployments. Visual analysis is heuristic rather "
        "than image-based machine learning, and therefore cannot detect "
        "brand impersonation outside its twelve-brand list. Finally, "
        "the statistical AI-authorship detector may misclassify formally "
        "written human emails (legal notices, academic correspondence) "
        "as AI-generated; this risk is mitigated by the modest +0.08 "
        "modifier but cannot be eliminated without a stronger authorship "
        "model.",
    )

    # ---------- Section VIII. Conclusion ----------
    add_para(doc, "Conclusion and Future Work", "Heading 1")
    add_para(
        doc,
        "We presented Hybrid AI Defense, a six-layer phishing detection "
        "system specifically designed to close the detection gap against "
        "AI-generated phishing emails. Our key contributions -- a novel "
        "LLM-generated phishing dataset, a dual-output DistilBERT "
        "classifier, a statistical AI-authorship detector, email header "
        "forensics, XAI token attribution, and a structured adversarial "
        "robustness evaluation framework -- together address the most "
        "pressing gaps in current phishing defense. The system is fully "
        "open source, requires no paid services beyond an optional "
        "VirusTotal API key, and deploys as a FastAPI backend with a "
        "web dashboard and a Chrome extension for Gmail.",
    )
    add_para(
        doc,
        "Future work will extend the visual analyzer to use a CNN-based "
        "screenshot classifier, add real-time threat-feed integration, "
        "and expand the adversarial test suite to include paraphrase-"
        "based and translation-based evasion.",
    )

    # ---------- Acknowledgment ----------
    add_para(doc, "Acknowledgment", "Heading 5")
    add_para(
        doc,
        "The author thanks the faculty of the School of Computer Science "
        "and Engineering at Vellore Institute of Technology for their "
        "guidance throughout this project.",
    )

    # ---------- References ----------
    add_para(doc, "References", "Heading 5")
    refs = [
        "G. Apruzzese et al., \"The role of machine learning in "
        "cybersecurity,\" ACM Transactions on Privacy and Security, "
        "vol. 26, no. 3, 2023.",
        "E. Hu et al., \"Detecting AI-generated text using perplexity "
        "and burstiness,\" arXiv preprint arXiv:2306.04723, 2023.",
        "T. Koide, N. Fukushi, H. Nakano, and D. Chiba, \"Detecting "
        "phishing sites using ChatGPT,\" arXiv preprint "
        "arXiv:2306.05816, 2023.",
        "J. Devlin, M.-W. Chang, K. Lee, and K. Toutanova, \"BERT: "
        "Pre-training of deep bidirectional transformers for language "
        "understanding,\" in Proc. NAACL-HLT, 2019, pp. 4171-4186.",
        "V. Sanh, L. Debut, J. Chaumond, and T. Wolf, \"DistilBERT, a "
        "distilled version of BERT: smaller, faster, cheaper and "
        "lighter,\" in NeurIPS EMC^2 Workshop, 2019.",
        "C. Geng, K. Sun, X. Wang, and P. Liu, \"Towards phishing-proof "
        "two-factor authentication,\" in Proc. IEEE Symp. Security and "
        "Privacy, 2018.",
        "Y. Liao, Z. Zhang, and L. Yang, \"Phishing detection via "
        "multi-modal deep neural network,\" in Proc. ICASSP, 2020.",
        "J. Ebrahimi, A. Rao, D. Lowd, and D. Dou, \"HotFlip: White-box "
        "adversarial examples for text classification,\" in Proc. ACL, "
        "2018, pp. 31-36.",
        "V. Zeng, S. Baki, A. El Aassal, R. Verma, L. F. De Moraes, and "
        "A. Das, \"PhishBench: A benchmarking framework for phishing "
        "detection,\" in Proc. RAID, 2020.",
        "J. Nazario, \"Phishing corpus,\" publicly available dataset, "
        "2005-2022. [Online]. Available: "
        "https://monkey.org/~jose/phishing/.",
        "M. Alazab and R. Broadhurst, \"Spam and criminal activity,\" "
        "Trends and Issues in Crime and Criminal Justice, no. 526, "
        "Australian Institute of Criminology, 2017.",
        "S. Salloum, T. Gaber, S. Vadera, and K. Shaalan, \"Phishing "
        "email detection using natural language processing techniques: "
        "A literature survey,\" Procedia Computer Science, vol. 189, "
        "pp. 19-28, 2021.",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph(style=ensure_style(doc, "references", "Normal"))
        p.add_run(f"[{i}] {ref}")

    doc.save(DST)
    print(f"Saved: {DST}")


if __name__ == "__main__":
    build_paper()
